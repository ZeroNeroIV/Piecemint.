"""Render invoice content for each export format. Client row comes from the DB model."""

from __future__ import annotations

import base64
import io
import os
import re
import tempfile

import httpx
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from invoice_document_render import document_subtotal, line_amount
from schemas import InvoiceExportConfig, PartyDetailsPayload

LOGO_MAX_BYTES = 2_000_000


def hex_to_rgb_01(hex_str: str) -> tuple[float, float, float]:
    h = hex_str.lstrip("#")
    return tuple(int(h[i : i + 2], 16) / 255.0 for i in (0, 2, 4))  # type: ignore[return-value]


def hex_to_rgb_int(hex_str: str) -> tuple[int, int, int]:
    h = hex_str.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _excel_font(family: str) -> str:
    return {
        "Helvetica": "Arial",
        "Helvetica-Bold": "Arial",
        "Times-Roman": "Times New Roman",
        "Times-Bold": "Times New Roman",
        "Courier": "Courier New",
        "Courier-Bold": "Courier New",
    }.get(family, "Arial")


def fetch_logo_bytes(url: str | None) -> bytes | None:
    if not url or not str(url).strip():
        return None
    try:
        u = str(url).strip()
        if not (u.startswith("http://") or u.startswith("https://")):
            return None
        r = httpx.get(u, follow_redirects=True, timeout=10.0)
        r.raise_for_status()
        if len(r.content) > LOGO_MAX_BYTES:
            return None
        return r.content
    except Exception:
        return None


def decode_data_url_logo(data_url: str | None) -> bytes | None:
    if not data_url or not str(data_url).strip().startswith("data:"):
        return None
    s = str(data_url).strip()
    m = re.match(r"^data:image/[^;]+;base64,(.+)$", s, re.DOTALL)
    if not m:
        return None
    try:
        raw = base64.b64decode(m.group(1), validate=True)
        if len(raw) > LOGO_MAX_BYTES:
            return None
        return raw
    except Exception:
        return None


def resolve_logo_bytes(cfg: InvoiceExportConfig) -> bytes | None:
    data_url = getattr(cfg, "logo_data_url", None)
    if data_url:
        b = decode_data_url_logo(data_url)
        if b:
            return b
    return fetch_logo_bytes(getattr(cfg, "logo_url", None))


def _format_party_lines(
    party: PartyDetailsPayload, fallback_name: str, fallback_email: str
) -> list[str]:
    lines: list[str] = []
    n = (party.legal_name or "").strip() or fallback_name
    if n:
        lines.append(n)
    if (party.address or "").strip():
        for part in party.address.split("\n")[:8]:
            if part.strip():
                lines.append(part.strip()[:120])
    if (party.tax_id or "").strip():
        lines.append(f"Tax ID: {party.tax_id.strip()}")
    em = (party.email or "").strip() or fallback_email
    if em:
        lines.append(f"Email: {em}")
    if (party.phone or "").strip():
        lines.append(f"Phone: {party.phone.strip()}")
    return lines if lines else ["—"]


def _build_pdf_legacy(
    name: str,
    email: str,
    client_id: str,
    total_billed: float,
    cfg: InvoiceExportConfig,
) -> bytes:
    pr, pg, pb = hex_to_rgb_01(cfg.primary_color)
    sr, sg, sb = hex_to_rgb_01(cfg.secondary_color)
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    logo_data = resolve_logo_bytes(cfg)
    y_cursor = height - 72

    if logo_data:
        try:
            ir = ImageReader(io.BytesIO(logo_data))
            lw, lh = 72, 72
            p.drawImage(ir, 72, y_cursor - lh, width=lw, height=lh, preserveAspectRatio=True, mask="auto")
        except Exception:
            pass
        y_cursor -= 84
    else:
        y_cursor -= 12

    p.setFont(cfg.font_family, min(cfg.font_size + 8, 24))
    p.setFillColorRGB(pr, pg, pb)
    p.drawString(72, y_cursor, "INVOICE")
    y_cursor -= 8
    p.setStrokeColorRGB(sr, sg, sb)
    p.setLineWidth(1.2)
    p.line(72, y_cursor, width - 72, y_cursor)
    y_cursor -= 28

    p.setFont(cfg.font_family, cfg.font_size)
    p.setFillColorRGB(pr, pg, pb)
    p.drawString(72, y_cursor, f"Bill to: {name}")
    y_cursor -= 18
    p.drawString(72, y_cursor, f"Client ID: {client_id}")
    y_cursor -= 18
    p.drawString(72, y_cursor, f"Email: {email}")
    y_cursor -= 18
    p.setFillColorRGB(sr, sg, sb)
    p.setFont(cfg.font_family, min(cfg.font_size + 1, 20))
    p.drawString(72, y_cursor, f"Total billed: ${total_billed:,.2f}")
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer.read()


def _build_pdf_with_document(
    name: str,
    email: str,
    client_id: str,
    total_billed: float,
    cfg: InvoiceExportConfig,
) -> bytes:
    d = cfg.invoice_document
    if d is None:
        return _build_pdf_legacy(name, email, client_id, total_billed, cfg)

    pr, pg, pb = hex_to_rgb_01(cfg.primary_color)
    sr, sg, sb = hex_to_rgb_01(cfg.secondary_color)
    sub = document_subtotal(d, total_billed)

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    logo_data = resolve_logo_bytes(cfg)
    y = height - 72

    if logo_data:
        try:
            ir = ImageReader(io.BytesIO(logo_data))
            p.drawImage(ir, 72, y - 72, width=72, height=72, preserveAspectRatio=True, mask="auto")
        except Exception:
            pass
        y -= 84
    else:
        y -= 12

    p.setFont(cfg.font_family, min(cfg.font_size + 8, 24))
    p.setFillColorRGB(pr, pg, pb)
    p.drawString(72, y, "INVOICE")
    y -= 10
    p.setStrokeColorRGB(sr, sg, sb)
    p.setLineWidth(1.0)
    p.line(72, y, width - 72, y)
    y -= 22

    fs = cfg.font_size
    p.setFont(cfg.font_family, fs)
    p.setFillColorRGB(pr, pg, pb)
    inv = (d.invoice_number or "").strip() or "—"
    issue = (d.issue_date or "").strip() or "—"
    due = (d.due_date or "").strip() or "—"
    for lab, val in [
        ("Invoice number", inv),
        ("Issue date", issue),
        ("Due date", due),
        ("Client ID (system)", client_id),
    ]:
        p.drawString(72, y, f"{lab}: {val}")
        y -= 16
    y -= 6

    def block(title: str, party: PartyDetailsPayload, cust: bool) -> None:
        nonlocal y
        p.setFont(cfg.font_family, fs + 1)
        p.setFillColorRGB(pr, pg, pb)
        p.drawString(72, y, title)
        y -= 16
        p.setFont(cfg.font_family, fs)
        for line in _format_party_lines(party, name if cust else "", email if cust else ""):
            p.drawString(72, y, line[:100])
            y -= 13
        y -= 4

    block("Customer (bill to)", d.customer, True)
    block("Supplier (your business)", d.supplier, False)

    p.setFont(cfg.font_family, fs + 1)
    p.setFillColorRGB(pr, pg, pb)
    p.drawString(72, y, "Line items")
    y -= 16
    p.setFont(cfg.font_family, max(fs - 1, 7))
    p.drawString(72, y, "Description")
    p.drawString(250, y, "Qty")
    p.drawString(300, y, "Unit")
    p.drawString(360, y, "Rate")
    p.drawString(420, y, "Amount")
    y -= 12
    p.setStrokeColorRGB(sr, sg, sb)
    p.line(72, y, width - 72, y)
    y -= 14

    p.setFont(cfg.font_family, fs)
    p.setFillColorRGB(pr, pg, pb)
    for li in d.line_items:
        if y < 100:
            p.showPage()
            y = height - 72
        p.drawString(72, y, (li.description or "—")[:45])
        p.drawString(250, y, f"{float(li.quantity or 0):g}")
        p.drawString(300, y, (li.unit_label or "units")[:10])
        p.drawString(360, y, f"${float(li.unit_price or 0):,.2f}")
        p.drawString(420, y, f"${line_amount(li):,.2f}")
        y -= 14

    y -= 6
    p.setFont(cfg.font_family, min(fs + 2, 20))
    p.setFillColorRGB(sr, sg, sb)
    p.drawString(72, y, f"Amount due: ${sub:,.2f}")
    y -= 22

    if (d.notes or "").strip():
        p.setFont(cfg.font_family, fs)
        p.setFillColorRGB(pr, pg, pb)
        p.drawString(72, y, "Notes")
        y -= 14
        p.setFont(cfg.font_family, max(fs - 1, 8))
        for part in (d.notes or "").split("\n")[:18]:
            p.drawString(72, y, (part or "")[:110])
            y -= 12

    tr = cfg.tax_residency
    if tr and tr.country_code:
        y -= 4
        p.setFont(cfg.font_family, fs)
        p.setFillColorRGB(pr, pg, pb)
        p.drawString(72, y, f"Tax residency: {tr.country_code}")
        y -= 12
        p.setFont(cfg.font_family, max(fs - 1, 8))
        for k, v in (tr.additional_data or {}).items():
            p.drawString(72, y, f"  {k}: {v[:80]}")
            y -= 12

    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer.read()


def build_pdf(
    name: str,
    email: str,
    client_id: str,
    total_billed: float,
    cfg: InvoiceExportConfig,
) -> bytes:
    if cfg.invoice_document is None:
        return _build_pdf_legacy(name, email, client_id, total_billed, cfg)
    return _build_pdf_with_document(name, email, client_id, total_billed, cfg)


def build_xlsx(
    name: str,
    email: str,
    client_id: str,
    total_billed: float,
    cfg: InvoiceExportConfig,
) -> bytes:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as SheetImage
    from openpyxl.styles import Font

    pr, pg, pb = hex_to_rgb_int(cfg.primary_color)
    sr, sg, sb = hex_to_rgb_int(cfg.secondary_color)
    prgb = f"FF{pr:02X}{pg:02X}{pb:02X}"
    srgb = f"FF{sr:02X}{sg:02X}{sb:02X}"
    font_name = _excel_font(cfg.font_family)

    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"
    if ws is None:
        raise RuntimeError("No active sheet")

    row = 1
    logo_data = resolve_logo_bytes(cfg)
    # openpyxl reads the image file from disk when wb.save() runs — keep the file until after save
    logo_temp_path: str | None = None
    logo_cell_a_width: float | None = None
    if logo_data:
        tmp_created: str | None = None
        try:
            from PIL import Image  # type: ignore[import-untyped]

            im = Image.open(io.BytesIO(logo_data))
            if im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGB")
            # Shrink to fit a single on-sheet cell; keep aspect ratio
            im.thumbnail((120, 120), Image.Resampling.LANCZOS)
            tw, th = im.size
            fd, tmp_created = tempfile.mkstemp(suffix=".png")
            os.close(fd)
            im.save(tmp_created, "PNG")
            sh = SheetImage(tmp_created)
            sh.width = tw
            sh.height = th
            # Size cell A1 to contain the image (width ≈ char units, height in points; ~96 dpi)
            # col width: ~7 px per character for default 11pt Calibri
            col_w = max(10.0, min(30.0, (tw / 7.0) + 0.85))
            row_h = max(15.0, (th * 72.0) / 96.0)
            ws.row_dimensions[1].height = row_h
            ws.column_dimensions["A"].width = col_w
            logo_cell_a_width = col_w
            # Anchor top-left in cell A1 (drawing is sized to match the cell box)
            ws.add_image(sh, "A1")
            logo_temp_path = tmp_created
            row = 2
        except Exception:
            if tmp_created and os.path.isfile(tmp_created):
                try:
                    os.unlink(tmp_created)
                except OSError:
                    pass
            logo_temp_path = None

    ws[f"A{row}"] = "INVOICE"
    ws[f"A{row}"].font = Font(
        name=font_name, size=cfg.font_size + 6, bold=True, color=prgb
    )
    row += 2

    if cfg.invoice_document is not None:
        d = cfg.invoice_document
        sub = document_subtotal(d, total_billed)
        for label, value in [
            ("Invoice number", (d.invoice_number or "").strip() or "—"),
            ("Issue date", (d.issue_date or "").strip() or "—"),
            ("Due date", (d.due_date or "").strip() or "—"),
            ("Client ID (system)", client_id),
        ]:
            c1 = ws.cell(row=row, column=1, value=label)
            c1.font = Font(name=font_name, size=cfg.font_size, color=prgb, bold=True)
            c2 = ws.cell(row=row, column=2, value=value)
            c2.font = Font(name=font_name, size=cfg.font_size, color=srgb)
            row += 1
        row += 1
        c1 = ws.cell(row=row, column=1, value="Customer (bill to)")
        c1.font = Font(name=font_name, size=cfg.font_size, color=prgb, bold=True)
        row += 1
        for pline in _format_party_lines(d.customer, name, email):
            ws.cell(row=row, column=1, value=pline)
            row += 1
        row += 1
        c1 = ws.cell(row=row, column=1, value="Supplier")
        c1.font = Font(name=font_name, size=cfg.font_size, color=prgb, bold=True)
        row += 1
        for pline in _format_party_lines(d.supplier, "", ""):
            ws.cell(row=row, column=1, value=pline)
            row += 1
        row += 1
        headers = ["Description", "Qty", "Unit", "Rate", "Line amount"]
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=row, column=col, value=h)
            c.font = Font(name=font_name, size=cfg.font_size, color=prgb, bold=True)
        row += 1
        for li in d.line_items:
            ws.cell(row=row, column=1, value=(li.description or "—")[:200])
            ws.cell(row=row, column=2, value=float(li.quantity or 0))
            ws.cell(row=row, column=3, value=(li.unit_label or "units")[:20])
            ws.cell(row=row, column=4, value=float(li.unit_price or 0))
            ws.cell(row=row, column=5, value=round(line_amount(li), 2))
            for c in range(1, 6):
                ws.cell(row=row, column=c).font = Font(
                    name=font_name, size=cfg.font_size, color=srgb
                )
            row += 1
        row += 1
        c1 = ws.cell(row=row, column=1, value="Amount due")
        c1.font = Font(name=font_name, size=cfg.font_size + 1, color=prgb, bold=True)
        c2 = ws.cell(row=row, column=2, value=f"${sub:,.2f}")
        c2.font = Font(name=font_name, size=cfg.font_size + 1, color=srgb, bold=True)
        row += 1
        if (d.notes or "").strip():
            row += 1
            c1 = ws.cell(row=row, column=1, value="Notes")
            c1.font = Font(name=font_name, size=cfg.font_size, color=prgb, bold=True)
            row += 1
            for nline in (d.notes or "").split("\n")[:20]:
                ws.cell(row=row, column=1, value=(nline or "")[:500])
                row += 1
        tr = cfg.tax_residency
        if tr and tr.country_code:
            row += 1
            ws.cell(row=row, column=1, value=f"Tax residency: {tr.country_code}")
            row += 1
            for k, v in (tr.additional_data or {}).items():
                ws.cell(row=row, column=1, value=f"{k}: {v}")
                row += 1
    else:
        for label, value in [
            ("Bill to", name),
            ("Client ID", client_id),
            ("Email", email),
            ("Total billed", f"${total_billed:,.2f}"),
        ]:
            c1 = ws.cell(row=row, column=1, value=label)
            c1.font = Font(name=font_name, size=cfg.font_size, color=prgb, bold=True)
            c2 = ws.cell(row=row, column=2, value=value)
            c2.font = Font(name=font_name, size=cfg.font_size, color=srgb)
            row += 1

    # Wide enough for label column; do not shrink below logo cell width if set
    ws.column_dimensions["A"].width = (
        max(20.0, logo_cell_a_width) if logo_cell_a_width is not None else 20.0
    )
    ws.column_dimensions["B"].width = 40

    buf = io.BytesIO()
    try:
        wb.save(buf)
    finally:
        if logo_temp_path and os.path.isfile(logo_temp_path):
            try:
                os.unlink(logo_temp_path)
            except OSError:
                pass
    buf.seek(0)
    return buf.read()


def build_docx(
    name: str,
    email: str,
    client_id: str,
    total_billed: float,
    cfg: InvoiceExportConfig,
) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    pr, pg, pb = hex_to_rgb_int(cfg.primary_color)
    sr, sg, sb = hex_to_rgb_int(cfg.secondary_color)

    doc = Document()
    logo_data = resolve_logo_bytes(cfg)
    if logo_data:
        try:
            from PIL import Image  # type: ignore[import-untyped]

            im = Image.open(io.BytesIO(logo_data))
            fd, tmpp = tempfile.mkstemp(suffix=".png")
            os.close(fd)
            try:
                im.save(tmpp, "PNG")
                doc.add_picture(tmpp, width=Inches(1.5))
            finally:
                os.unlink(tmpp)
        except Exception:
            pass

    h = doc.add_heading("INVOICE", level=0)
    for r in h.runs:
        r.font.size = Pt(min(cfg.font_size + 6, 28))
        r.font.name = "Helvetica" if "Helvetica" in cfg.font_family else "Times New Roman" if "Times" in cfg.font_family else "Courier New"
        r.font.color.rgb = RGBColor(pr, pg, pb)

    if cfg.invoice_document is not None:
        d = cfg.invoice_document
        sub = document_subtotal(d, total_billed)
        for label, value in [
            ("Invoice number", (d.invoice_number or "").strip() or "—"),
            ("Issue date", (d.issue_date or "").strip() or "—"),
            ("Due date", (d.due_date or "").strip() or "—"),
            ("Client ID (system)", client_id),
        ]:
            pp = doc.add_paragraph()
            pp.add_run(f"{label}: ").bold = True
            rr = pp.add_run(value)
            rr.font.size = Pt(cfg.font_size)
            rr.font.color.rgb = RGBColor(pr, pg, pb)
        h2 = doc.add_paragraph()
        h2.add_run("Customer (bill to)").bold = True
        h2.runs[0].font.size = Pt(cfg.font_size)
        h2.runs[0].font.color.rgb = RGBColor(pr, pg, pb)
        for pline in _format_party_lines(d.customer, name, email):
            doc.add_paragraph(pline)
        h3 = doc.add_paragraph()
        h3.add_run("Supplier").bold = True
        h3.runs[0].font.size = Pt(cfg.font_size)
        h3.runs[0].font.color.rgb = RGBColor(pr, pg, pb)
        for pline in _format_party_lines(d.supplier, "", ""):
            doc.add_paragraph(pline)
        pli = doc.add_paragraph()
        rli = pli.add_run("Line items")
        rli.bold = True
        rli.font.size = Pt(cfg.font_size)
        rli.font.color.rgb = RGBColor(pr, pg, pb)
        t = doc.add_paragraph("Description  |  Qty  |  Unit  |  Rate  |  Amount")
        for li in d.line_items:
            doc.add_paragraph(
                f"{(li.description or '—')[:200]}  |  {float(li.quantity or 0)}  |  {li.unit_label}  |  ${float(li.unit_price or 0):,.2f}  |  ${line_amount(li):,.2f}"
            )
        ptot = doc.add_paragraph()
        ptot.add_run("Amount due: ").bold = True
        rtot = ptot.add_run(f"${sub:,.2f}")
        rtot.font.size = Pt(cfg.font_size + 1)
        rtot.font.color.rgb = RGBColor(sr, sg, sb)
        if (d.notes or "").strip():
            pn = doc.add_paragraph()
            rn = pn.add_run("Notes")
            rn.bold = True
            rn.font.size = Pt(cfg.font_size)
            doc.add_paragraph((d.notes or "")[:4000])
        tr = cfg.tax_residency
        if tr and tr.country_code:
            doc.add_paragraph(f"Tax residency: {tr.country_code}")
            for k, v in (tr.additional_data or {}).items():
                doc.add_paragraph(f"{k}: {v}")
    else:
        p1 = doc.add_paragraph()
        p1.add_run("Bill to: ").bold = True
        run = p1.add_run(name)
        run.font.size = Pt(cfg.font_size)
        run.font.color.rgb = RGBColor(pr, pg, pb)

        p2 = doc.add_paragraph()
        p2.add_run("Client ID: ").bold = True
        r2 = p2.add_run(client_id)
        r2.font.size = Pt(cfg.font_size)
        r2.font.color.rgb = RGBColor(pr, pg, pb)

        p3 = doc.add_paragraph()
        p3.add_run("Email: ").bold = True
        r3 = p3.add_run(email)
        r3.font.size = Pt(cfg.font_size)
        r3.font.color.rgb = RGBColor(pr, pg, pb)

        p4 = doc.add_paragraph()
        p4.add_run("Total billed: ").bold = True
        r4 = p4.add_run(f"${total_billed:,.2f}")
        r4.font.size = Pt(cfg.font_size + 1)
        r4.font.color.rgb = RGBColor(sr, sg, sb)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render_invoice(
    name: str,
    email: str,
    client_id: str,
    total_billed: float,
    cfg: InvoiceExportConfig,
) -> tuple[bytes, str, str]:
    fmt = cfg.output_format
    if fmt == "pdf":
        return build_pdf(name, email, client_id, total_billed, cfg), "application/pdf", ".pdf"
    if fmt == "xlsx":
        return build_xlsx(name, email, client_id, total_billed, cfg), (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ), ".xlsx"
    if fmt == "docx":
        return build_docx(name, email, client_id, total_billed, cfg), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ), ".docx"
    raise ValueError("Unsupported output_format")
